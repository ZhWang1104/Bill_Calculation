"""
================================================================================
    Bill Parser GUI Application
    Author: Zh Wang
    License: MIT

    Description:
        该程序提供一个图形界面，从当前目录下读取 .txt 账单文件，
        解析条目、计算总金额，并将结果导出为 Word (.docx) 和 Excel (.xlsx)。
        同时在界面中展示总行数、有效行数、错误行与累计总金额。

    (C) 2024 Zh Wang. All rights reserved.
    This is an original work. Distribution and modification permitted with attribution.
================================================================================
"""

import os
import re
import sys
import tkinter as tk
from tkinter import messagebox, Text, Frame, Button
from tkinter.font import Font
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font as XLFont, Alignment as XLAlignment

# ----------------- 简单的配色 / UI 常量 -----------------
BG_COLOR = "#f4f6f7"       # 整体背景
PANEL_COLOR = "#ffffff"    # 卡片背景
ACCENT_COLOR = "#2c3e50"   # 深色文字
BUTTON_BG = "#3498db"      # 按钮背景
BUTTON_BG_ACTIVE = "#2980b9"
BUTTON_FG = "#ffffff"

# Windows 高 DPI 处理
if sys.platform == "win32":
    try:
        import ctypes
        ctypes.windll.shcore.SetProcessDpiAwareness(1)
    except Exception:
        try:
            ctypes.windll.user32.SetProcessDPIAware()
        except Exception:
            pass


# --------------------------------------------
# 核心处理函数：解析 txt 数据、生成 Word & Excel
# --------------------------------------------
def process_txt_file(file_path, output_folder, txt_name, doc_title):
    docx_path = os.path.join(output_folder, txt_name.replace(".txt", ".docx"))
    doc = Document()

    # 1. Word 文档标题（大写、14pt、居中加粗）
    title = doc.add_heading('', level=1)
    run = title.add_run(doc_title.upper())
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. 创建三列表格：Name / Quantity / Price
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Quantity'
    hdr_cells[2].text = 'Price'

    # 为 Excel 准备的行缓存（只记录分类行和商品行）
    excel_rows = []

    # 正则：完整“名称 数量 $价格” 或 “名称 $价格”
    pattern_full = re.compile(r"(.+?)\s+([\d.]+(?:kg|pcs|set|p)?\.?)\s*\$([\d.]+)")
    pattern_short = re.compile(r"(.+?)\s*\$([\d.]+)")

    with open(file_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    total_lines = 0          # 文件中的所有行计数
    valid_lines = 0          # 成功解析的商品行计数
    grand_total = 0.0        # 全部商品累加总价
    error_lines = []         # 无法解析的行列表

    # 小标题：Word 里插入一行合并 3 列、居中加粗；Excel 里记录一条 'category'
    def insert_category_title(name: str):
        row = table.add_row()
        cell = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
        p = cell.paragraphs[0]
        r = p.add_run(name)
        r.bold = True
        r.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        excel_rows.append(('category', name))

    # 开始逐行处理
    for idx, line in enumerate(lines, 1):
        total_lines += 1
        text = line.strip()
        if not text:
            continue

        # 如果一行出现两个及以上 '$'，直接报错，不参与计算
        if text.count('$') > 1:
            error_lines.append(f"[Line {idx}] ❌ 含有多个 $，无法处理：{text}")
            continue

        m_full = pattern_full.match(text)
        m_short = pattern_short.match(text)

        # === 小标题判断规则 ===
        # 条件：不是 full/short 商品行，且整行不含 '$'
        # 例如 "8nummmmmm" 这种没有价格的行，也按小标题处理
        if not m_full and not m_short and ('$' not in text):
            insert_category_title(text)
            continue

        # 普通商品行
        if m_full:
            name, qty, price = m_full.groups()
        elif m_short:
            name, price = m_short.groups()
            qty = ""
        else:
            # 剩下：包含 '$' 但格式不合法，视为错误行
            error_lines.append(f"[Line {idx}] ❌ 无法解析：{text}")
            continue

        # Word 表格中插入商品行
        row_cells = table.add_row().cells
        row_cells[0].text = name.strip()
        row_cells[1].text = qty.strip()
        row_cells[2].text = f"${price.strip()}"

        # 记录到 Excel 缓存
        excel_rows.append(('item', name.strip(), qty.strip(), price.strip()))

        # 金额累加
        try:
            price_val = float(price)
            grand_total += price_val
            valid_lines += 1
        except Exception:
            error_lines.append(f"[Line {idx}] ⚠️ 价格解析失败：{price}")

    # Word 中插入最终的 Grand Total（合并三列、12pt、加粗居中）
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    p = cell.paragraphs[0]
    r = p.add_run(f"Grand Total: ${grand_total:.2f}")
    r.bold = True
    r.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加 Word 文末 Footer（左对齐、10pt）
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_run = footer.add_run(
        "This bill was made by Sophia\nGenerated on: " +
        datetime.now().strftime("%B %d, %Y")
    )
    footer_run.font.size = Pt(10)

    # 保存 Word
    doc.save(docx_path)

    # --------- 写入 Excel ---------
    excel_path = os.path.join(output_folder, txt_name.replace('.txt', '.xlsx'))
    wb = Workbook()
    ws = wb.active
    try:
        ws.title = doc_title[:31]
    except Exception:
        pass

    # 列宽调整
    ws.column_dimensions['A'].width = 30  # 名称
    ws.column_dimensions['B'].width = 15  # 数量
    ws.column_dimensions['C'].width = 15  # 价格

    # 第 1 行：1×3 大表头，内容和 Word 一样
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=3)
    title_cell = ws.cell(row=1, column=1, value=doc_title.upper())
    title_cell.font = XLFont(name="Times New Roman", bold=True, size=14)
    title_cell.alignment = XLAlignment(horizontal='center', vertical='center')

    # 第 2 行：列名表头
    header_row_idx = 2
    headers = ['Name', 'Quantity', 'Price']
    for col, h in enumerate(headers, start=1):
        c = ws.cell(row=header_row_idx, column=col, value=h)
        c.font = XLFont(name="Times New Roman", bold=True)
        c.alignment = XLAlignment(horizontal='center', vertical='center')

    # 数据从第 3 行开始
    row_idx = header_row_idx + 1

    # 写入缓存行
    for entry in excel_rows:
        typ = entry[0]
        if typ == 'category':
            name = entry[1]
            ws.merge_cells(start_row=row_idx, start_column=1, end_row=row_idx, end_column=3)
            c = ws.cell(row=row_idx, column=1, value=name)
            c.font = XLFont(name="Times New Roman", bold=True)
            c.alignment = XLAlignment(horizontal='center', vertical='center')
            row_idx += 1

        elif typ == 'item':
            _, name, qty, price_str = entry
            c1 = ws.cell(row=row_idx, column=1, value=name)
            c2 = ws.cell(row=row_idx, column=2, value=qty)
            c3 = ws.cell(row=row_idx, column=3, value=f"${price_str}")
            # 账单中的数据居中，字体 Times New Roman
            for c in (c1, c2, c3):
                c.font = XLFont(name="Times New Roman")
                c.alignment = XLAlignment(horizontal='center', vertical='center')
            row_idx += 1

    # 添加 Grand Total（2×3 合并，加粗居中）
    start_gt_row = row_idx
    end_gt_row = row_idx + 1
    ws.merge_cells(start_row=start_gt_row, start_column=1, end_row=end_gt_row, end_column=3)
    gt_cell = ws.cell(row=start_gt_row, column=1, value=f"Grand Total: ${grand_total:.2f}")
    gt_cell.font = XLFont(name="Times New Roman", bold=True, size=12)
    gt_cell.alignment = XLAlignment(horizontal='center', vertical='center')
    row_idx = end_gt_row + 1

    # 添加落款信息（2×3 合并，加粗左对齐 + 手动增大行高保证显示完整）
    footer_text = (
        "Invoice prepared by Sophia\n"
        f"Invoice date: {datetime.now().strftime('%d %B %Y')}\n"
        "If you have any queries regarding this invoice, please contact us.\n"
        "Many thanks for your custom."
    )
    start_footer_row = row_idx
    end_footer_row = row_idx + 1  # 2 行 × 3 列
    ws.merge_cells(
        start_row=start_footer_row,
        start_column=1,
        end_row=end_footer_row,
        end_column=3
    )
    footer_cell = ws.cell(row=start_footer_row, column=1, value=footer_text)
    footer_cell.font = XLFont(name="Times New Roman", bold=True)
    footer_cell.alignment = XLAlignment(
        horizontal='left',
        vertical='top',   # 从上往下排
        wrap_text=True
    )
    # 手动拉高这两行的行高，避免内容被截断
    for r in range(start_footer_row, end_footer_row + 1):
        ws.row_dimensions[r].height = 30

    # 保存 Excel
    wb.save(excel_path)

    return error_lines, total_lines, valid_lines, grand_total, docx_path, excel_path


# --------------------------------------------
# 自定义输入对话框：输入 .txt 文件名（居中弹出）
# --------------------------------------------
def ask_txt_filename():
    """
    自定义的模态对话框，要求用户输入 .txt 文件名（不需要后缀），
    返回值为用户输入的纯文件名（无 .txt）或 None（取消）。
    """
    dlg = tk.Toplevel()
    dlg.title("选择账单文件")
    dlg.resizable(False, False)
    dlg.configure(bg=BG_COLOR)

    # 对话框大小
    width, height = 400, 180

    # 计算屏幕居中位置
    screen_w = dlg.winfo_screenwidth()
    screen_h = dlg.winfo_screenheight()
    x = (screen_w - width) // 2
    y = (screen_h - height) // 2
    dlg.geometry(f"{width}x{height}+{x}+{y}")

    # 外层留一点空白
    container = tk.Frame(dlg, bg=PANEL_COLOR, bd=1, relief="solid")
    container.pack(fill="both", expand=True, padx=12, pady=12)

    title_lbl = tk.Label(
        container,
        text="账单文件选择",
        bg=PANEL_COLOR,
        fg=ACCENT_COLOR,
        font=("Segoe UI", 11, "bold")
    )
    title_lbl.pack(pady=(10, 4))

    hint_lbl = tk.Label(
        container,
        text="请输入 .txt 文件名（例如：0603，对应 0603.txt）",
        bg=PANEL_COLOR,
        fg=ACCENT_COLOR,
        font=("Segoe UI", 9)
    )
    hint_lbl.pack(pady=(0, 6))

    entry = tk.Entry(container, width=28, font=("Segoe UI", 10))
    entry.pack(pady=(0, 10))
    entry.focus_set()

    result = {'value': None}

    def on_ok():
        val = entry.get().strip()
        if val == "":
            result['value'] = None
        else:
            result['value'] = val
        dlg.destroy()

    def on_cancel():
        result['value'] = None
        dlg.destroy()

    btn_frame = tk.Frame(container, bg=PANEL_COLOR)
    btn_frame.pack(pady=(5, 12))

    ok_btn = tk.Button(
        btn_frame, text="确认", width=10,
        command=on_ok,
        bg=BUTTON_BG, fg=BUTTON_FG,
        activebackground=BUTTON_BG_ACTIVE,
        activeforeground=BUTTON_FG,
        relief="flat", bd=0
    )
    ok_btn.pack(side="left", padx=8)

    cancel_btn = tk.Button(
        btn_frame, text="取消", width=10,
        command=on_cancel,
        bg="#bdc3c7", fg="#2c3e50",
        activebackground="#95a5a6",
        activeforeground="#2c3e50",
        relief="flat", bd=0
    )
    cancel_btn.pack(side="right", padx=8)

    dlg.bind("<Return>", lambda e: on_ok())
    dlg.bind("<Escape>", lambda e: on_cancel())

    dlg.grab_set()
    dlg.wait_window()
    return result['value']


# --------------------------------------------
# 自定义输入对话框：输入账单标题（居中弹出）
# --------------------------------------------
def ask_bill_title():
    """
    模态对话框，让用户输入账单标题（可留空，留空时使用默认 "PRODUCTS"）。
    返回字符串或 None。
    """
    dlg = tk.Toplevel()
    dlg.title("账单标题")
    dlg.resizable(False, False)
    dlg.configure(bg=BG_COLOR)

    width, height = 360, 170
    screen_w = dlg.winfo_screenwidth()
    screen_h = dlg.winfo_screenheight()
    x = (screen_w - width) // 2
    y = (screen_h - height) // 2
    dlg.geometry(f"{width}x{height}+{x}+{y}")

    container = tk.Frame(dlg, bg=PANEL_COLOR, bd=1, relief="solid")
    container.pack(fill="both", expand=True, padx=12, pady=12)

    title_lbl = tk.Label(
        container,
        text="设置账单标题",
        bg=PANEL_COLOR,
        fg=ACCENT_COLOR,
        font=("Segoe UI", 11, "bold")
    )
    title_lbl.pack(pady=(10, 4))

    lbl = tk.Label(
        container,
        text="请输入账单标题（默认：PRODUCTS）",
        bg=PANEL_COLOR,
        fg=ACCENT_COLOR,
        font=("Segoe UI", 9)
    )
    lbl.pack(pady=(0, 6))

    entry = tk.Entry(container, width=28, font=("Segoe UI", 10))
    entry.pack(pady=(0, 10))
    entry.focus_set()

    result = {'value': None}

    def on_ok():
        result['value'] = entry.get().strip()
        dlg.destroy()

    def on_cancel():
        result['value'] = None
        dlg.destroy()

    btn_frame = tk.Frame(container, bg=PANEL_COLOR)
    btn_frame.pack(pady=(5, 12))

    ok_btn = tk.Button(
        btn_frame, text="确认", width=10,
        command=on_ok,
        bg=BUTTON_BG, fg=BUTTON_FG,
        activebackground=BUTTON_BG_ACTIVE,
        activeforeground=BUTTON_FG,
        relief="flat", bd=0
    )
    ok_btn.pack(side="left", padx=8)

    cancel_btn = tk.Button(
        btn_frame, text="取消", width=10,
        command=on_cancel,
        bg="#bdc3c7", fg="#2c3e50",
        activebackground="#95a5a6",
        activeforeground="#2c3e50",
        relief="flat", bd=0
    )
    cancel_btn.pack(side="right", padx=8)

    dlg.bind("<Return>", lambda e: on_ok())
    dlg.bind("<Escape>", lambda e: on_cancel())

    dlg.grab_set()
    dlg.wait_window()
    return result['value']


# --------------------------------------------
# 结果显示窗口
# --------------------------------------------
def show_result_window(errors, total, valid, amount, docx_path, excel_path):
    def on_close():
        # 退出程序
        result_win.destroy()
        sys.exit()

    def on_continue():
        # 关闭结果窗口，重新进入主循环处理下一个文件
        result_win.destroy()
        run_main_loop()

    result_win = tk.Tk()
    result_win.title("处理结果")
    result_win.resizable(False, False)

    # 窗口大小 & 居中
    width, height = 760, 520
    screen_w = result_win.winfo_screenwidth()
    screen_h = result_win.winfo_screenheight()
    x = (screen_w - width) // 2
    y = (screen_h - height) // 2
    result_win.geometry(f"{width}x{height}+{x}+{y}")

    result_win.protocol("WM_DELETE_WINDOW", on_close)
    result_win.configure(bg="#f5f5f5")

    # ===== 顶部标题 =====
    title_lbl = tk.Label(
        result_win,
        text="本次账单处理结果",
        bg="#f5f5f5",
        fg="#333333",
        font=("Microsoft YaHei", 14, "bold")
    )
    title_lbl.pack(pady=(12, 6))

    # ===== 概要信息区域 =====
    summary_frame = tk.Frame(result_win, bg="#f5f5f5")
    summary_frame.pack(fill="x", padx=20, pady=(0, 8))

    summary_text = [
        f"Word 文件 ：{docx_path}",
        f"Excel 文件：{excel_path}",
        f"总行数      ：{total}",
        f"有效处理行数：{valid}",
        f"累计金额 (Grand Total)：${amount:.2f}",
    ]
    summary_label = tk.Label(
        summary_frame,
        text="\n".join(summary_text),
        bg="#f5f5f5",
        fg="#333333",
        font=("Microsoft YaHei", 11),
        justify="left",
        anchor="w"
    )
    summary_label.pack(fill="x")

    # 分割线
    sep = tk.Frame(result_win, bg="#d0d0d0", height=1)
    sep.pack(fill="x", padx=20, pady=(4, 6))

    # ===== 错误信息区域 =====
    center_frame = tk.Frame(result_win, bg="#f5f5f5")
    center_frame.pack(fill="both", expand=True, padx=20, pady=(0, 8))

    if errors:
        err_title = tk.Label(
            center_frame,
            text="以下行为无法处理，请检查：",
            bg="#f5f5f5",
            fg="#aa0000",
            font=("Microsoft YaHei", 11, "bold"),
            anchor="w",
            justify="left"
        )
        err_title.pack(fill="x", pady=(0, 4))

        text_frame = tk.Frame(center_frame, bg="#f5f5f5")
        text_frame.pack(fill="both", expand=True)

        scrollbar = tk.Scrollbar(text_frame)
        scrollbar.pack(side="right", fill="y")

        err_text = Text(
            text_frame,
            width=90,
            height=14,
            font=("Consolas", 11),
            padx=8,
            pady=6,
            wrap="word"
        )
        err_text.pack(side="left", fill="both", expand=True)

        err_text.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=err_text.yview)

        for err in errors:
            err_text.insert("end", err + "\n")
        err_text.config(state="disabled")
    else:
        ok_label = tk.Label(
            center_frame,
            text="所有数据处理成功。",
            bg="#f5f5f5",
            fg="#006400",
            font=("Microsoft YaHei", 12, "bold"),
            anchor="w",
            justify="left"
        )
        ok_label.pack(fill="x", pady=(6, 0))

    # ===== 底部按钮区域（关键：再处理 / 退出）=====
    btn_frame = tk.Frame(result_win, bg="#f5f5f5")
    btn_frame.pack(pady=(6, 12))

    btn_style = {
        "font": ("Microsoft YaHei", 11),
        "width": 18,
        "relief": "flat",
        "bd": 0,
    }

    btn_continue = tk.Button(
        btn_frame,
        text="处理其他文件",
        command=on_continue,
        bg="#4a90e2",
        fg="white",
        activebackground="#357ab8",
        activeforeground="white",
        **btn_style
    )
    btn_continue.pack(side="left", padx=20)

    btn_exit = tk.Button(
        btn_frame,
        text="退出程序",
        command=on_close,
        bg="#cccccc",
        fg="#333333",
        activebackground="#b3b3b3",
        activeforeground="#333333",
        **btn_style
    )
    btn_exit.pack(side="right", padx=20)

    result_win.mainloop()



# --------------------------------------------
# 主循环：请求输入、调度处理、显示结果
# --------------------------------------------
def run_main_loop():
    root = tk.Tk()
    root.withdraw()

    while True:
        # 1) 输入 .txt 文件名（不带后缀）
        txt_input = ask_txt_filename()
        if txt_input is None:
            sys.exit()

        if not txt_input.lower().endswith(".txt"):
            txt_input += ".txt"

        # 2) 从可执行/脚本所在目录查找文件
        if getattr(sys, 'frozen', False):
            search_path = os.path.dirname(sys.executable)
        else:
            search_path = os.path.dirname(__file__)
        file_path = os.path.join(search_path, txt_input)

        if not os.path.isfile(file_path):
            messagebox.showerror("错误", f"文件不存在：{file_path}")
            continue

        # 3) 输入账单标题
        title_input = ask_bill_title()
        if not title_input:
            title_input = "PRODUCTS"

        # 4) 解析并生成 Word / Excel
        errors, total, valid, amount, docx_path, excel_path = process_txt_file(
            file_path, search_path, txt_input, title_input
        )

        # 5) 显示结果
        show_result_window(errors, total, valid, amount, docx_path, excel_path)
        break


# --------------------------------------------
# 程序入口
# --------------------------------------------
if __name__ == "__main__":
    run_main_loop()
